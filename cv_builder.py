"""
cv_builder.py — Mesin pembangun dokumen Curriculum Vitae.

Menghasilkan dua format keluaran:
  * PDF  (reportlab)  — format utama, lengkap dengan halaman lampiran
  * DOCX (python-docx) — format tambahan untuk keperluan penyuntingan

Standar dokumen formal Indonesia:
  Kertas A4 potret, margin atas 2,5 cm / bawah 2,5 cm / kiri 3 cm / kanan 2 cm,
  huruf Times New Roman 12 pt untuk isi dan 14 pt tebal untuk judul.
"""

import io
from dataclasses import dataclass, field
from datetime import date
from typing import List, Optional, Dict, Any

from pypdf import PdfReader, PdfWriter

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage,
    KeepTogether,
)

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import utils
from utils import (
    NAMA_KEGIATAN, INSTANSI_KEMENTERIAN, INSTANSI_UNIVERSITAS,
    INSTANSI_FAKULTAS, INSTANSI_ALAMAT, LOGO_PATH, logo_tersedia,
    format_tanggal_indonesia, aman,
)

# =====================================================================
# TETAPAN GAYA
# =====================================================================

# Palet warna sesuai standar aplikasi: biru laut & hijau karang
WARNA_UTAMA = colors.HexColor("#1B3F6B")   # Biru laut
WARNA_AKSEN = colors.HexColor("#2D7D46")   # Hijau karang
WARNA_TEKS = colors.HexColor("#1A1A1A")
WARNA_GARIS = colors.HexColor("#7A8CA0")
WARNA_LATAR_LABEL = colors.HexColor("#EAF0F6")

DOCX_WARNA_UTAMA = RGBColor(0x1B, 0x3F, 0x6B)
DOCX_WARNA_AKSEN = RGBColor(0x2D, 0x7D, 0x46)
DOCX_LATAR_LABEL = "EAF0F6"

FONT = "Times-Roman"
FONT_TEBAL = "Times-Bold"
FONT_MIRING = "Times-Italic"
FONT_DOCX = "Times New Roman"

# Margin standar dokumen formal Indonesia
MARGIN_ATAS = 2.5 * cm
MARGIN_BAWAH = 2.5 * cm
MARGIN_KIRI = 3.0 * cm
MARGIN_KANAN = 2.0 * cm

# Lebar area kerja = 21 cm − 3 cm − 2 cm = 16 cm
LEBAR_KERJA = 16.0 * cm

JENIS_TENAGA_AHLI = "Tenaga Ahli"
JENIS_PENYELAM = "Tenaga Spesialis Penyelaman"


# =====================================================================
# STRUKTUR DATA
# =====================================================================

@dataclass
class Lampiran:
    """Satu berkas lampiran yang akan disisipkan di akhir dokumen CV."""
    judul: str                      # contoh: "KTP"
    nama_berkas: str                # nama berkas asli hasil unggahan
    data: bytes                     # isi berkas
    adalah_pdf: bool = False        # True bila berkas berformat PDF


@dataclass
class CVData:
    """Seluruh data yang dibutuhkan untuk menyusun satu dokumen CV."""

    # --- Jenis dokumen ---
    jenis_cv: str = JENIS_TENAGA_AHLI

    # --- Data pribadi ---
    nama: str = ""
    tempat_lahir: str = ""
    tanggal_lahir: Optional[date] = None
    jenis_kelamin: str = ""
    agama: str = ""
    kewarganegaraan: str = "Indonesia"
    alamat: str = ""
    telepon: str = ""
    email: str = ""
    nomor_ktp: str = ""
    npwp: str = ""

    # --- Data kepegawaian (opsional, untuk ASN) ---
    nip: str = ""
    nidn: str = ""
    jabatan: str = ""
    afiliasi: str = "Fakultas Perikanan dan Ilmu Kelautan, Universitas Sam Ratulangi"

    # --- Peran dalam tim swakelola ---
    peran_tim: str = ""
    peran_teknis: str = ""

    # --- Tabel isian ---
    pendidikan: List[Dict[str, Any]] = field(default_factory=list)
    sertifikasi: List[Dict[str, Any]] = field(default_factory=list)
    pengalaman_kerja: List[Dict[str, Any]] = field(default_factory=list)
    publikasi: List[Dict[str, Any]] = field(default_factory=list)
    bidang_keahlian: List[str] = field(default_factory=list)

    # --- Khusus Tenaga Spesialis Penyelaman ---
    lisensi_jenis: str = ""
    lisensi_nomor: str = ""
    lisensi_level: str = ""
    lisensi_terbit: Optional[date] = None
    lisensi_berlaku: Optional[date] = None

    medis_nomor: str = ""
    medis_penerbit: str = ""
    medis_terbit: Optional[date] = None
    medis_berlaku: Optional[date] = None

    total_jam_selam: int = 0
    keahlian_selam: List[str] = field(default_factory=list)
    pengalaman_selam: List[Dict[str, Any]] = field(default_factory=list)

    # --- Pernyataan dan tanda tangan ---
    tempat_ttd: str = "Bitung"
    tanggal_ttd: Optional[date] = None
    nama_terang: str = ""

    # --- Ringkasan keterkaitan dengan kegiatan (dapat diisi otomatis/manual) ---
    ringkasan_afiliasi: str = ""

    # --- Berkas ---
    foto: Optional[bytes] = None
    lampiran: List[Lampiran] = field(default_factory=list)

    def ttl(self) -> str:
        """Menggabungkan tempat dan tanggal lahir menjadi satu baris teks."""
        tgl = format_tanggal_indonesia(self.tanggal_lahir)
        if self.tempat_lahir and tgl:
            return f"{self.tempat_lahir}, {tgl}"
        return aman(self.tempat_lahir or tgl)

    def adalah_penyelam(self) -> bool:
        """Mengembalikan True bila dokumen berjenis Spesialis Penyelaman."""
        return self.jenis_cv == JENIS_PENYELAM


PERNYATAAN_BAKU = (
    "Demikian daftar riwayat hidup ini saya buat dengan sebenar-benarnya "
    "untuk dapat dipergunakan sebagaimana mestinya."
)


# =====================================================================
# BAGIAN 1 — PEMBANGUN PDF
# =====================================================================

def _gaya_pdf() -> Dict[str, ParagraphStyle]:
    """Menyusun kumpulan gaya paragraf untuk dokumen PDF."""
    return {
        "kop_besar": ParagraphStyle(
            "kop_besar", fontName=FONT_TEBAL, fontSize=13, leading=15,
            alignment=TA_CENTER, textColor=WARNA_TEKS,
        ),
        "kop_kecil": ParagraphStyle(
            "kop_kecil", fontName=FONT_TEBAL, fontSize=11, leading=13,
            alignment=TA_CENTER, textColor=WARNA_TEKS,
        ),
        "kop_kementerian": ParagraphStyle(
            "kop_kementerian", fontName=FONT_TEBAL, fontSize=9.5, leading=11,
            alignment=TA_CENTER, textColor=WARNA_TEKS,
        ),
        "kop_alamat": ParagraphStyle(
            "kop_alamat", fontName=FONT, fontSize=8, leading=9.5,
            alignment=TA_CENTER, textColor=WARNA_TEKS,
        ),
        "judul": ParagraphStyle(
            "judul", fontName=FONT_TEBAL, fontSize=14, leading=17,
            alignment=TA_CENTER, textColor=WARNA_UTAMA, spaceAfter=2,
        ),
        "kegiatan": ParagraphStyle(
            "kegiatan", fontName=FONT_TEBAL, fontSize=11, leading=13,
            alignment=TA_CENTER, textColor=WARNA_AKSEN,
        ),
        "seksi": ParagraphStyle(
            "seksi", fontName=FONT_TEBAL, fontSize=12, leading=14,
            textColor=WARNA_UTAMA, spaceBefore=10, spaceAfter=4,
        ),
        "isi": ParagraphStyle(
            "isi", fontName=FONT, fontSize=12, leading=15,
            textColor=WARNA_TEKS, alignment=TA_LEFT,
        ),
        "isi_kecil": ParagraphStyle(
            "isi_kecil", fontName=FONT, fontSize=10.5, leading=13,
            textColor=WARNA_TEKS, alignment=TA_LEFT,
        ),
        "isi_kecil_tengah": ParagraphStyle(
            "isi_kecil_tengah", fontName=FONT, fontSize=10.5, leading=13,
            textColor=WARNA_TEKS, alignment=TA_CENTER,
        ),
        "label": ParagraphStyle(
            "label", fontName=FONT_TEBAL, fontSize=12, leading=15,
            textColor=WARNA_UTAMA,
        ),
        "kepala_tabel": ParagraphStyle(
            "kepala_tabel", fontName=FONT_TEBAL, fontSize=10.5, leading=13,
            textColor=colors.white, alignment=TA_CENTER,
        ),
        "pernyataan": ParagraphStyle(
            "pernyataan", fontName=FONT, fontSize=12, leading=18,
            textColor=WARNA_TEKS, alignment=TA_JUSTIFY, spaceBefore=6,
        ),
        "ttd": ParagraphStyle(
            "ttd", fontName=FONT, fontSize=12, leading=16,
            textColor=WARNA_TEKS, alignment=TA_CENTER,
        ),
        "ttd_nama": ParagraphStyle(
            "ttd_nama", fontName=FONT_TEBAL, fontSize=12, leading=16,
            textColor=WARNA_TEKS, alignment=TA_CENTER,
        ),
        "lampiran_judul": ParagraphStyle(
            "lampiran_judul", fontName=FONT_TEBAL, fontSize=14, leading=18,
            alignment=TA_CENTER, textColor=WARNA_UTAMA,
        ),
    }


def _kop_pdf(g: Dict[str, ParagraphStyle], data: CVData) -> List:
    """Menyusun kop resmi instansi beserta pas foto di pojok kanan atas."""
    elemen = []

    # Kolom kiri: logo instansi
    if logo_tersedia():
        try:
            sel_logo = RLImage(LOGO_PATH, width=2.0 * cm, height=2.0 * cm)
        except Exception:
            sel_logo = Paragraph("[LOGO]", g["isi_kecil_tengah"])
    else:
        sel_logo = Paragraph("[LOGO]", g["isi_kecil_tengah"])

    # Kolom tengah: identitas instansi
    sel_teks = [
        Paragraph(INSTANSI_KEMENTERIAN, g["kop_kementerian"]),
        Paragraph(INSTANSI_UNIVERSITAS, g["kop_besar"]),
        Paragraph(INSTANSI_FAKULTAS, g["kop_kecil"]),
        Paragraph(INSTANSI_ALAMAT, g["kop_alamat"]),
    ]

    # Kolom kanan: pas foto 3x4
    if data.foto:
        try:
            foto_bersih = utils.potong_pas_foto(data.foto)
            sel_foto = RLImage(
                io.BytesIO(foto_bersih), width=2.25 * cm, height=3.0 * cm
            )
        except Exception:
            sel_foto = Paragraph("[PAS FOTO<br/>3 x 4]", g["isi_kecil_tengah"])
    else:
        sel_foto = Paragraph("[PAS FOTO<br/>3 x 4]", g["isi_kecil_tengah"])

    tabel = Table(
        [[sel_logo, sel_teks, sel_foto]],
        colWidths=[2.4 * cm, 10.8 * cm, 2.8 * cm],
    )
    tabel.setStyle(TableStyle([
        ("VALIGN", (0, 0), (1, 0), "MIDDLE"),
        ("VALIGN", (2, 0), (2, 0), "TOP"),
        ("ALIGN", (0, 0), (0, 0), "CENTER"),
        ("ALIGN", (2, 0), (2, 0), "RIGHT"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        # Bingkai pas foto agar tetap terlihat meskipun foto belum diunggah
        ("BOX", (2, 0), (2, 0), 0.75, WARNA_GARIS),
    ]))
    elemen.append(tabel)

    # Garis tebal pemisah kop
    garis = Table([[""]], colWidths=[LEBAR_KERJA], rowHeights=[2])
    garis.setStyle(TableStyle([
        ("LINEBELOW", (0, 0), (-1, -1), 2, colors.black),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
    ]))
    elemen.append(Spacer(1, 4))
    elemen.append(garis)
    elemen.append(Spacer(1, 12))

    # Judul dokumen dan nama kegiatan
    elemen.append(Paragraph("CURRICULUM VITAE", g["judul"]))
    elemen.append(Paragraph(data.jenis_cv.upper(), g["kegiatan"]))
    elemen.append(Spacer(1, 4))
    elemen.append(Paragraph(NAMA_KEGIATAN, g["isi_kecil_tengah"]))
    elemen.append(Spacer(1, 8))
    return elemen


def _judul_seksi_pdf(g: Dict[str, ParagraphStyle], teks: str) -> Table:
    """Membuat judul bagian dengan garis bawah berwarna aksen."""
    tabel = Table([[Paragraph(teks.upper(), g["seksi"])]], colWidths=[LEBAR_KERJA])
    tabel.setStyle(TableStyle([
        ("LINEBELOW", (0, 0), (-1, -1), 1, WARNA_AKSEN),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
    ]))
    return tabel


def _tabel_label_pdf(g: Dict[str, ParagraphStyle],
                     baris: List[tuple]) -> Table:
    """Membuat tabel dua kolom bergaya label–nilai untuk data pribadi."""
    isi = []
    for label, nilai in baris:
        isi.append([
            Paragraph(label, g["label"]),
            Paragraph(f": {aman(nilai)}", g["isi"]),
        ])

    tabel = Table(isi, colWidths=[5.2 * cm, 10.8 * cm])
    tabel.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BACKGROUND", (0, 0), (0, -1), WARNA_LATAR_LABEL),
        ("GRID", (0, 0), (-1, -1), 0.5, WARNA_GARIS),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return tabel


def _tabel_data_pdf(g: Dict[str, ParagraphStyle], kepala: List[str],
                    baris: List[List[str]], lebar: List[float]) -> Table:
    """Membuat tabel berkepala berwarna untuk data pendidikan, pengalaman, dll."""
    isi = [[Paragraph(k, g["kepala_tabel"]) for k in kepala]]
    for b in baris:
        isi.append([Paragraph(str(sel), g["isi_kecil"]) for sel in b])

    tabel = Table(isi, colWidths=lebar, repeatRows=1)
    tabel.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), WARNA_UTAMA),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.5, WARNA_GARIS),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.white, colors.HexColor("#F5F8FB")]),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return tabel


def _blok_ttd_pdf(g: Dict[str, ParagraphStyle], data: CVData) -> List:
    """Menyusun blok pernyataan dan kolom tanda tangan."""
    tgl = format_tanggal_indonesia(data.tanggal_ttd or date.today())
    nama_terang = data.nama_terang.strip() or data.nama

    blok = [
        Spacer(1, 10),
        Paragraph(PERNYATAAN_BAKU, g["pernyataan"]),
        Spacer(1, 14),
    ]

    isi_ttd = [
        [Paragraph(f"{aman(data.tempat_ttd)}, {tgl}", g["ttd"])],
        [Paragraph("Yang membuat pernyataan,", g["ttd"])],
        [Spacer(1, 62)],
        [Paragraph(f"<u>{nama_terang}</u>", g["ttd_nama"])],
    ]
    if data.nip.strip():
        isi_ttd.append([Paragraph(f"NIP. {data.nip}", g["ttd"])])

    tabel = Table(isi_ttd, colWidths=[7.0 * cm], hAlign="RIGHT")
    tabel.setStyle(TableStyle([
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
    ]))
    blok.append(KeepTogether(tabel))
    return blok


def _isi_pdf_tenaga_ahli(g: Dict[str, ParagraphStyle], data: CVData) -> List:
    """Menyusun bagian isi CV untuk jenis Tenaga Ahli."""
    e = []

    # A. Data pribadi
    e.append(_judul_seksi_pdf(g, "A. Data Pribadi"))
    e.append(_tabel_label_pdf(g, _baris_data_pribadi(data)))

    # B. Riwayat pendidikan
    e.append(_judul_seksi_pdf(g, "B. Riwayat Pendidikan"))
    baris = [
        [b.get("Jenjang", ""), b.get("Jurusan", ""), b.get("Universitas", ""),
         b.get("Kota", ""), b.get("Tahun Lulus", "")]
        for b in data.pendidikan
    ]
    if baris:
        e.append(_tabel_data_pdf(
            g, ["Jenjang", "Jurusan", "Universitas", "Kota", "Tahun"], baris,
            [2.0 * cm, 4.0 * cm, 5.0 * cm, 2.8 * cm, 2.2 * cm],
        ))
    else:
        e.append(Paragraph("Belum diisi.", g["isi"]))

    # C. Bidang keahlian
    e.append(_judul_seksi_pdf(g, "C. Bidang Keahlian"))
    if data.bidang_keahlian:
        teks = "  •  ".join(data.bidang_keahlian)
    else:
        teks = "-"
    e.append(Paragraph(teks, g["isi"]))

    # D. Sertifikasi kompetensi
    e.append(_judul_seksi_pdf(g, "D. Sertifikasi Kompetensi"))
    baris = [
        [b.get("Nama Sertifikat", ""), b.get("Nomor", ""),
         b.get("Lembaga Penerbit", ""), b.get("Tahun", ""),
         b.get("Masa Berlaku", "")]
        for b in data.sertifikasi
    ]
    if baris:
        e.append(_tabel_data_pdf(
            g, ["Nama Sertifikat", "Nomor", "Lembaga", "Tahun", "Berlaku"],
            baris, [4.4 * cm, 3.0 * cm, 4.0 * cm, 2.0 * cm, 2.6 * cm],
        ))
    else:
        e.append(Paragraph("Tidak ada.", g["isi"]))

    # E. Pengalaman kerja
    e.append(_judul_seksi_pdf(g, "E. Pengalaman Kerja"))
    baris = []
    for i, b in enumerate(data.pengalaman_kerja, start=1):
        baris.append([
            str(i), b.get("Nama Pekerjaan/Proyek", ""), b.get("Tahun", ""),
            b.get("Jabatan/Posisi", ""), b.get("Instansi Pemberi Kerja", ""),
            b.get("Lokasi", ""),
        ])
    if baris:
        e.append(_tabel_data_pdf(
            g, ["No", "Pekerjaan / Proyek", "Tahun", "Jabatan",
                "Instansi", "Lokasi"],
            baris,
            [1.0 * cm, 4.4 * cm, 2.2 * cm, 2.8 * cm, 3.2 * cm, 2.4 * cm],
        ))
    else:
        e.append(Paragraph("Belum diisi.", g["isi"]))

    # F. Publikasi ilmiah (opsional)
    if data.publikasi:
        e.append(_judul_seksi_pdf(g, "F. Publikasi Ilmiah"))
        baris = [
            [str(i), b.get("Judul", ""), b.get("Jurnal/Prosiding", ""),
             b.get("Tahun", "")]
            for i, b in enumerate(data.publikasi, start=1)
        ]
        e.append(_tabel_data_pdf(
            g, ["No", "Judul", "Jurnal / Prosiding", "Tahun"], baris,
            [1.0 * cm, 7.4 * cm, 5.4 * cm, 2.2 * cm],
        ))

    return e


def _isi_pdf_penyelam(g: Dict[str, ParagraphStyle], data: CVData) -> List:
    """Menyusun bagian isi CV untuk jenis Tenaga Spesialis Penyelaman."""
    e = []

    # A. Data pribadi
    e.append(_judul_seksi_pdf(g, "A. Data Pribadi"))
    e.append(_tabel_label_pdf(g, _baris_data_pribadi(data)))

    # B. Sertifikat selam utama
    e.append(_judul_seksi_pdf(g, "B. Sertifikat Selam Utama"))
    e.append(_tabel_label_pdf(g, [
        ("Jenis Lisensi", data.lisensi_jenis),
        ("Nomor Sertifikat", data.lisensi_nomor),
        ("Level Sertifikasi", data.lisensi_level),
        ("Tanggal Terbit", format_tanggal_indonesia(data.lisensi_terbit)),
        ("Masa Berlaku", format_tanggal_indonesia(data.lisensi_berlaku)
         or "Seumur hidup"),
    ]))

    # C. Sertifikat medis selam
    e.append(_judul_seksi_pdf(g, "C. Sertifikat Medis Selam (Fit to Dive)"))
    e.append(_tabel_label_pdf(g, [
        ("Nomor Sertifikat", data.medis_nomor),
        ("Dokter / Klinik Penerbit", data.medis_penerbit),
        ("Tanggal Terbit", format_tanggal_indonesia(data.medis_terbit)),
        ("Berlaku Sampai", format_tanggal_indonesia(data.medis_berlaku)),
    ]))

    # D. Kompetensi penyelaman
    e.append(_judul_seksi_pdf(g, "D. Kompetensi Penyelaman"))
    e.append(_tabel_label_pdf(g, [
        ("Total Jam Selam Terverifikasi", f"{data.total_jam_selam} jam"),
    ]))
    e.append(Spacer(1, 6))
    e.append(Paragraph("<b>Keahlian Khusus Penyelaman:</b>", g["isi"]))
    if data.keahlian_selam:
        for k in data.keahlian_selam:
            e.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&#9642; {k}", g["isi"]))
    else:
        e.append(Paragraph("&nbsp;&nbsp;&nbsp;-", g["isi"]))

    # E. Pengalaman penyelaman
    e.append(_judul_seksi_pdf(g, "E. Pengalaman Penyelaman"))
    baris = []
    for i, b in enumerate(data.pengalaman_selam, start=1):
        baris.append([
            str(i), b.get("Lokasi Penyelaman", ""), b.get("Tahun", ""),
            b.get("Jenis Kegiatan", ""), b.get("Kedalaman Maks (m)", ""),
            b.get("Lama Kegiatan", ""), b.get("Pemberi Kerja", ""),
        ])
    if baris:
        e.append(_tabel_data_pdf(
            g, ["No", "Lokasi", "Thn", "Jenis Kegiatan", "Maks (m)",
                "Lama", "Pemberi Kerja"],
            baris,
            [0.9 * cm, 3.2 * cm, 1.3 * cm, 3.4 * cm, 1.7 * cm,
             2.0 * cm, 3.5 * cm],
        ))
    else:
        e.append(Paragraph("Belum diisi.", g["isi"]))

    # F. Pendidikan formal
    e.append(_judul_seksi_pdf(g, "F. Pendidikan Formal"))
    baris = [
        [b.get("Jenjang", ""), b.get("Jurusan", ""), b.get("Universitas", ""),
         b.get("Kota", ""), b.get("Tahun Lulus", "")]
        for b in data.pendidikan
    ]
    if baris:
        e.append(_tabel_data_pdf(
            g, ["Jenjang", "Jurusan", "Universitas", "Kota", "Tahun"], baris,
            [2.0 * cm, 4.0 * cm, 5.0 * cm, 2.8 * cm, 2.2 * cm],
        ))
    else:
        e.append(Paragraph("Belum diisi.", g["isi"]))

    # G. Sertifikasi tambahan (opsional)
    if data.sertifikasi:
        e.append(_judul_seksi_pdf(g, "G. Sertifikasi Kompetensi Lainnya"))
        baris = [
            [b.get("Nama Sertifikat", ""), b.get("Nomor", ""),
             b.get("Lembaga Penerbit", ""), b.get("Tahun", ""),
             b.get("Masa Berlaku", "")]
            for b in data.sertifikasi
        ]
        e.append(_tabel_data_pdf(
            g, ["Nama Sertifikat", "Nomor", "Lembaga", "Tahun", "Berlaku"],
            baris, [4.4 * cm, 3.0 * cm, 4.0 * cm, 2.0 * cm, 2.6 * cm],
        ))

    return e


def _baris_data_pribadi(data: CVData) -> List[tuple]:
    """Menyusun daftar pasangan label–nilai untuk tabel data pribadi."""
    baris = [
        ("Nama Lengkap", data.nama),
        ("Tempat, Tanggal Lahir", data.ttl()),
        ("Jenis Kelamin", data.jenis_kelamin),
        ("Agama", data.agama),
        ("Kewarganegaraan", data.kewarganegaraan),
        ("Alamat Lengkap", data.alamat),
        ("Nomor Telepon", data.telepon),
        ("Surel (Email)", data.email),
        ("Nomor KTP (NIK)", data.nomor_ktp),
    ]
    if data.npwp.strip():
        baris.append(("NPWP", data.npwp))
    if data.nip.strip():
        baris.append(("NIP", data.nip))
    if data.nidn.strip():
        baris.append(("NIDN", data.nidn))
    if data.jabatan.strip():
        baris.append(("Jabatan Fungsional", data.jabatan))
    if data.afiliasi.strip():
        baris.append(("Afiliasi Institusi", data.afiliasi))
    if data.peran_tim.strip():
        peran = data.peran_tim
        if data.peran_teknis.strip():
            peran = f"{peran} — {data.peran_teknis}"
        baris.append(("Peran dalam Tim", peran))
    return baris


def _gambar_kaki_halaman(kanvas, dokumen):
    """Menggambar nomor halaman pada setiap lembar dokumen PDF."""
    kanvas.saveState()
    kanvas.setFont(FONT, 9)
    kanvas.setFillColor(colors.HexColor("#666666"))
    kanvas.drawCentredString(
        A4[0] / 2.0, 1.4 * cm, f"Halaman {kanvas.getPageNumber()}"
    )
    kanvas.restoreState()


def _pdf_utama(data: CVData) -> bytes:
    """Membangun halaman utama CV (tanpa lampiran) sebagai berkas PDF."""
    penyangga = io.BytesIO()
    dokumen = SimpleDocTemplate(
        penyangga,
        pagesize=A4,
        topMargin=MARGIN_ATAS,
        bottomMargin=MARGIN_BAWAH,
        leftMargin=MARGIN_KIRI,
        rightMargin=MARGIN_KANAN,
        title=f"Curriculum Vitae — {data.nama}",
        author=data.nama,
        subject=NAMA_KEGIATAN,
    )

    g = _gaya_pdf()
    elemen = _kop_pdf(g, data)

    if data.adalah_penyelam():
        elemen += _isi_pdf_penyelam(g, data)
    else:
        elemen += _isi_pdf_tenaga_ahli(g, data)

    if data.ringkasan_afiliasi.strip():
        elemen.append(_judul_seksi_pdf(g, "Ringkasan Keterkaitan dengan Kegiatan"))
        elemen.append(Paragraph(data.ringkasan_afiliasi, g["isi"]))

    elemen += _blok_ttd_pdf(g, data)

    dokumen.build(
        elemen,
        onFirstPage=_gambar_kaki_halaman,
        onLaterPages=_gambar_kaki_halaman,
    )
    penyangga.seek(0)
    return penyangga.getvalue()


def _halaman_lampiran_gambar(judul: str, data_gambar: bytes) -> bytes:
    """Membuat satu halaman PDF berisi judul lampiran dan satu gambar.

    Gambar diskalakan secara proporsional agar muat pada area kerja
    tanpa terpotong maupun mengalami distorsi rasio.
    """
    penyangga = io.BytesIO()
    kanvas = rl_canvas.Canvas(penyangga, pagesize=A4)
    lebar_hal, tinggi_hal = A4

    # Judul lampiran
    kanvas.setFont(FONT_TEBAL, 14)
    kanvas.setFillColor(WARNA_UTAMA)
    kanvas.drawCentredString(lebar_hal / 2.0, tinggi_hal - MARGIN_ATAS, judul)

    # Garis bawah judul
    kanvas.setStrokeColor(WARNA_AKSEN)
    kanvas.setLineWidth(1)
    y_garis = tinggi_hal - MARGIN_ATAS - 8
    kanvas.line(MARGIN_KIRI, y_garis, lebar_hal - MARGIN_KANAN, y_garis)

    # Area yang tersedia untuk gambar
    area_lebar = lebar_hal - MARGIN_KIRI - MARGIN_KANAN
    area_tinggi = y_garis - MARGIN_BAWAH - 20

    try:
        bersih = utils.normalisasi_gambar(data_gambar, maks_sisi=2000)
        pembaca = ImageReader(io.BytesIO(bersih))
        lebar_img, tinggi_img = pembaca.getSize()

        skala = min(area_lebar / lebar_img, area_tinggi / tinggi_img)
        lebar_akhir = lebar_img * skala
        tinggi_akhir = tinggi_img * skala

        x = (lebar_hal - lebar_akhir) / 2.0
        y = MARGIN_BAWAH + (area_tinggi - tinggi_akhir) / 2.0

        kanvas.drawImage(
            pembaca, x, y, width=lebar_akhir, height=tinggi_akhir,
            preserveAspectRatio=True, anchor="c",
        )
        # Bingkai tipis mengelilingi gambar
        kanvas.setStrokeColor(WARNA_GARIS)
        kanvas.setLineWidth(0.5)
        kanvas.rect(x, y, lebar_akhir, tinggi_akhir)
    except Exception:
        kanvas.setFont(FONT, 12)
        kanvas.setFillColor(colors.black)
        kanvas.drawCentredString(
            lebar_hal / 2.0, tinggi_hal / 2.0,
            "Gambar lampiran tidak dapat ditampilkan.",
        )

    kanvas.showPage()
    kanvas.save()
    penyangga.seek(0)
    return penyangga.getvalue()


def _halaman_judul_lampiran(judul: str, nama_berkas: str) -> bytes:
    """Membuat halaman pemisah bertuliskan judul lampiran berformat PDF."""
    penyangga = io.BytesIO()
    kanvas = rl_canvas.Canvas(penyangga, pagesize=A4)
    lebar_hal, tinggi_hal = A4

    kanvas.setFont(FONT_TEBAL, 14)
    kanvas.setFillColor(WARNA_UTAMA)
    kanvas.drawCentredString(lebar_hal / 2.0, tinggi_hal - MARGIN_ATAS, judul)

    kanvas.setStrokeColor(WARNA_AKSEN)
    kanvas.setLineWidth(1)
    y_garis = tinggi_hal - MARGIN_ATAS - 8
    kanvas.line(MARGIN_KIRI, y_garis, lebar_hal - MARGIN_KANAN, y_garis)

    kanvas.setFont(FONT_MIRING, 11)
    kanvas.setFillColor(colors.HexColor("#555555"))
    kanvas.drawCentredString(
        lebar_hal / 2.0, tinggi_hal / 2.0,
        f"Berkas terlampir: {nama_berkas}",
    )
    kanvas.drawCentredString(
        lebar_hal / 2.0, tinggi_hal / 2.0 - 18,
        "Isi lampiran ditampilkan pada halaman berikutnya.",
    )

    kanvas.showPage()
    kanvas.save()
    penyangga.seek(0)
    return penyangga.getvalue()


def _rakit_dengan_lampiran(pdf_utama: bytes,
                           lampiran: List[Lampiran]) -> bytes:
    """Menggabungkan halaman utama CV dengan seluruh halaman lampiran.

    Lampiran gambar dirender menjadi satu halaman penuh; lampiran PDF
    diberi halaman pemisah lalu seluruh halaman aslinya disisipkan.
    """
    penulis = PdfWriter()

    for halaman in PdfReader(io.BytesIO(pdf_utama)).pages:
        penulis.add_page(halaman)

    for nomor, lam in enumerate(lampiran, start=1):
        judul = f"Lampiran {nomor} — {lam.judul}"
        try:
            if lam.adalah_pdf:
                pemisah = _halaman_judul_lampiran(judul, lam.nama_berkas)
                for halaman in PdfReader(io.BytesIO(pemisah)).pages:
                    penulis.add_page(halaman)
                for halaman in PdfReader(io.BytesIO(lam.data)).pages:
                    penulis.add_page(halaman)
            else:
                satu = _halaman_lampiran_gambar(judul, lam.data)
                for halaman in PdfReader(io.BytesIO(satu)).pages:
                    penulis.add_page(halaman)
        except Exception:
            # Lampiran rusak tetap diberi halaman keterangan agar tidak hilang
            pemisah = _halaman_judul_lampiran(
                judul, f"{lam.nama_berkas} (berkas tidak dapat dibaca)"
            )
            for halaman in PdfReader(io.BytesIO(pemisah)).pages:
                penulis.add_page(halaman)

    keluaran = io.BytesIO()
    penulis.write(keluaran)
    keluaran.seek(0)
    return keluaran.getvalue()


def generate_cv_pdf(data: CVData) -> bytes:
    """Menghasilkan dokumen CV lengkap berformat PDF beserta lampirannya.

    Ini adalah fungsi keluaran utama yang dipanggil oleh antarmuka aplikasi.
    """
    utama = _pdf_utama(data)
    if not data.lampiran:
        return utama
    return _rakit_dengan_lampiran(utama, data.lampiran)


# =====================================================================
# BAGIAN 2 — PEMBANGUN DOCX
# =====================================================================

def _atur_arsir_sel(sel, warna_heks: str):
    """Memberi warna latar pada sebuah sel tabel DOCX."""
    arsir = OxmlElement("w:shd")
    arsir.set(qn("w:val"), "clear")
    arsir.set(qn("w:fill"), warna_heks)
    sel._tc.get_or_add_tcPr().append(arsir)


def _hapus_tepi_sel(sel):
    """Menghilangkan seluruh garis tepi pada sebuah sel tabel DOCX."""
    tcPr = sel._tc.get_or_add_tcPr()
    tepi = OxmlElement("w:tcBorders")
    for sisi in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{sisi}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tepi.append(el)
    tcPr.append(tepi)


def _tulis(paragraf, teks: str, ukuran: int = 12, tebal: bool = False,
           miring: bool = False, warna: Optional[RGBColor] = None,
           garis_bawah: bool = False):
    """Menambahkan satu potongan teks berformat ke dalam paragraf DOCX."""
    r = paragraf.add_run(teks)
    r.font.name = FONT_DOCX
    r.font.size = Pt(ukuran)
    r.bold = tebal
    r.italic = miring
    r.underline = garis_bawah
    if warna is not None:
        r.font.color.rgb = warna
    # Pastikan huruf Times New Roman juga berlaku untuk karakter non-latin
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_DOCX)
    rFonts.set(qn("w:cs"), FONT_DOCX)
    return r


def _kop_docx(dok: Document, data: CVData):
    """Menyusun kop resmi instansi dan pas foto pada dokumen DOCX."""
    tabel = dok.add_table(rows=1, cols=3)
    tabel.alignment = WD_TABLE_ALIGNMENT.CENTER

    sel_logo, sel_teks, sel_foto = tabel.rows[0].cells
    sel_logo.width = Cm(2.4)
    sel_teks.width = Cm(10.8)
    sel_foto.width = Cm(2.8)

    # Logo instansi
    p_logo = sel_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo_tersedia():
        try:
            p_logo.add_run().add_picture(LOGO_PATH, width=Cm(2.0))
        except Exception:
            _tulis(p_logo, "[LOGO]", ukuran=9)
    else:
        _tulis(p_logo, "[LOGO]", ukuran=9)

    # Identitas instansi
    baris_kop = [
        (INSTANSI_KEMENTERIAN, 11, True),
        (INSTANSI_UNIVERSITAS, 13, True),
        (INSTANSI_FAKULTAS, 11, True),
        (INSTANSI_ALAMAT, 8, False),
    ]
    pertama = True
    for teks, ukuran, tebal in baris_kop:
        p = sel_teks.paragraphs[0] if pertama else sel_teks.add_paragraph()
        pertama = False
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _tulis(p, teks, ukuran=ukuran, tebal=tebal)

    # Pas foto 3x4
    p_foto = sel_foto.paragraphs[0]
    p_foto.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if data.foto:
        try:
            foto_bersih = utils.potong_pas_foto(data.foto)
            p_foto.add_run().add_picture(
                io.BytesIO(foto_bersih), width=Cm(2.25), height=Cm(3.0)
            )
        except Exception:
            _tulis(p_foto, "[PAS FOTO 3x4]", ukuran=9)
    else:
        _tulis(p_foto, "[PAS FOTO 3x4]", ukuran=9)

    for sel in tabel.rows[0].cells:
        _hapus_tepi_sel(sel)

    # Garis tebal pemisah kop
    p_garis = dok.add_paragraph()
    p_garis.paragraph_format.space_before = Pt(2)
    p_garis.paragraph_format.space_after = Pt(8)
    pPr = p_garis._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bawah = OxmlElement("w:bottom")
    bawah.set(qn("w:val"), "single")
    bawah.set(qn("w:sz"), "18")
    bawah.set(qn("w:space"), "1")
    bawah.set(qn("w:color"), "000000")
    pBdr.append(bawah)
    pPr.append(pBdr)

    # Judul dokumen
    p_judul = dok.add_paragraph()
    p_judul.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_judul.paragraph_format.space_after = Pt(0)
    _tulis(p_judul, "CURRICULUM VITAE", ukuran=14, tebal=True,
           warna=DOCX_WARNA_UTAMA)

    p_jenis = dok.add_paragraph()
    p_jenis.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_jenis.paragraph_format.space_after = Pt(2)
    _tulis(p_jenis, data.jenis_cv.upper(), ukuran=11, tebal=True,
           warna=DOCX_WARNA_AKSEN)

    p_keg = dok.add_paragraph()
    p_keg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_keg.paragraph_format.space_after = Pt(8)
    _tulis(p_keg, NAMA_KEGIATAN, ukuran=10)


def _judul_seksi_docx(dok: Document, teks: str):
    """Menambahkan judul bagian bergaris bawah pada dokumen DOCX."""
    p = dok.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    _tulis(p, teks.upper(), ukuran=12, tebal=True, warna=DOCX_WARNA_UTAMA)

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bawah = OxmlElement("w:bottom")
    bawah.set(qn("w:val"), "single")
    bawah.set(qn("w:sz"), "8")
    bawah.set(qn("w:space"), "2")
    bawah.set(qn("w:color"), "2D7D46")
    pBdr.append(bawah)
    pPr.append(pBdr)
    return p


def _tabel_label_docx(dok: Document, baris: List[tuple]):
    """Menambahkan tabel label–nilai dua kolom pada dokumen DOCX."""
    tabel = dok.add_table(rows=len(baris), cols=2)
    tabel.style = "Table Grid"
    for i, (label, nilai) in enumerate(baris):
        sel_kiri = tabel.cell(i, 0)
        sel_kanan = tabel.cell(i, 1)
        sel_kiri.width = Cm(5.2)
        sel_kanan.width = Cm(10.8)
        _atur_arsir_sel(sel_kiri, DOCX_LATAR_LABEL)
        _tulis(sel_kiri.paragraphs[0], label, tebal=True,
               warna=DOCX_WARNA_UTAMA)
        _tulis(sel_kanan.paragraphs[0], f": {aman(nilai)}")
    return tabel


def _tabel_data_docx(dok: Document, kepala: List[str],
                     baris: List[List[str]], lebar: List[float]):
    """Menambahkan tabel berkepala pada dokumen DOCX."""
    tabel = dok.add_table(rows=len(baris) + 1, cols=len(kepala))
    tabel.style = "Table Grid"

    for j, judul in enumerate(kepala):
        sel = tabel.cell(0, j)
        sel.width = Cm(lebar[j])
        _atur_arsir_sel(sel, "1B3F6B")
        p = sel.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _tulis(p, judul, ukuran=10, tebal=True,
               warna=RGBColor(0xFF, 0xFF, 0xFF))

    for i, isi_baris in enumerate(baris, start=1):
        for j, nilai in enumerate(isi_baris):
            sel = tabel.cell(i, j)
            sel.width = Cm(lebar[j])
            _tulis(sel.paragraphs[0], str(nilai), ukuran=10)
    return tabel


def _blok_ttd_docx(dok: Document, data: CVData):
    """Menambahkan blok pernyataan dan kolom tanda tangan pada DOCX."""
    p = dok.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    _tulis(p, PERNYATAAN_BAKU)

    tgl = format_tanggal_indonesia(data.tanggal_ttd or date.today())
    nama_terang = data.nama_terang.strip() or data.nama

    tabel = dok.add_table(rows=1, cols=2)
    tabel.alignment = WD_TABLE_ALIGNMENT.RIGHT
    sel_kosong, sel_ttd = tabel.rows[0].cells
    sel_kosong.width = Cm(8.5)
    sel_ttd.width = Cm(7.5)
    _hapus_tepi_sel(sel_kosong)
    _hapus_tepi_sel(sel_ttd)

    def _baris_ttd(teks, tebal=False, garis_bawah=False, pertama=False):
        p = sel_ttd.paragraphs[0] if pertama else sel_ttd.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        _tulis(p, teks, tebal=tebal, garis_bawah=garis_bawah)

    _baris_ttd(f"{aman(data.tempat_ttd)}, {tgl}", pertama=True)
    _baris_ttd("Yang membuat pernyataan,")
    for _ in range(4):
        _baris_ttd("")
    _baris_ttd(nama_terang, tebal=True, garis_bawah=True)
    if data.nip.strip():
        _baris_ttd(f"NIP. {data.nip}")


def _lampiran_docx(dok: Document, data: CVData):
    """Menambahkan halaman-halaman lampiran pada dokumen DOCX.

    Lampiran gambar disisipkan langsung sebagai gambar satu halaman penuh.
    Lampiran berformat PDF tidak dapat disisipkan ke dalam berkas Word,
    sehingga hanya dituliskan keterangan rujukan; versi PDF tetap memuat
    seluruh halaman lampiran secara utuh.
    """
    from docx.enum.text import WD_BREAK

    for nomor, lam in enumerate(data.lampiran, start=1):
        p_putus = dok.add_paragraph()
        p_putus.add_run().add_break(WD_BREAK.PAGE)

        p_judul = dok.add_paragraph()
        p_judul.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_judul.paragraph_format.space_after = Pt(10)
        _tulis(p_judul, f"Lampiran {nomor} — {lam.judul}", ukuran=14,
               tebal=True, warna=DOCX_WARNA_UTAMA)

        p_isi = dok.add_paragraph()
        p_isi.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if lam.adalah_pdf:
            _tulis(p_isi, f"Berkas terlampir: {lam.nama_berkas}", miring=True)
            p_ket = dok.add_paragraph()
            p_ket.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _tulis(p_ket,
                   "(Lampiran berformat PDF hanya dapat ditampilkan utuh "
                   "pada berkas keluaran PDF)", ukuran=10, miring=True)
        else:
            try:
                bersih = utils.normalisasi_gambar(lam.data, maks_sisi=1600)
                lebar_px, tinggi_px = utils.dimensi_gambar(bersih)
                # Skalakan agar muat pada area kerja 16 cm x 20 cm
                lebar_cm, tinggi_cm = 16.0, 20.0
                if lebar_px and tinggi_px:
                    skala = min(lebar_cm / lebar_px, tinggi_cm / tinggi_px)
                    lebar_akhir = lebar_px * skala
                else:
                    lebar_akhir = 14.0
                p_isi.add_run().add_picture(
                    io.BytesIO(bersih), width=Cm(min(lebar_akhir, 16.0))
                )
            except Exception:
                _tulis(p_isi,
                       f"Gambar lampiran {lam.nama_berkas} tidak dapat "
                       f"ditampilkan.", miring=True)


def generate_cv_docx(data: CVData) -> bytes:
    """Menghasilkan dokumen CV berformat DOCX beserta lampirannya.

    Format DOCX disediakan sebagai keluaran tambahan agar dokumen masih
    dapat disunting apabila diperlukan koreksi administratif.
    """
    dok = Document()

    # Margin standar dokumen formal Indonesia
    for bagian in dok.sections:
        bagian.top_margin = Cm(2.5)
        bagian.bottom_margin = Cm(2.5)
        bagian.left_margin = Cm(3.0)
        bagian.right_margin = Cm(2.0)

    # Huruf baku dokumen
    gaya_normal = dok.styles["Normal"]
    gaya_normal.font.name = FONT_DOCX
    gaya_normal.font.size = Pt(12)

    _kop_docx(dok, data)

    if data.adalah_penyelam():
        _isi_docx_penyelam(dok, data)
    else:
        _isi_docx_tenaga_ahli(dok, data)

    if data.ringkasan_afiliasi.strip():
        _judul_seksi_docx(dok, "Ringkasan Keterkaitan dengan Kegiatan")
        p_ringkasan = dok.add_paragraph()
        p_ringkasan.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_ringkasan.paragraph_format.line_spacing = 1.5
        _tulis(p_ringkasan, data.ringkasan_afiliasi)

    _blok_ttd_docx(dok, data)

    if data.lampiran:
        _lampiran_docx(dok, data)

    keluaran = io.BytesIO()
    dok.save(keluaran)
    keluaran.seek(0)
    return keluaran.getvalue()


def _isi_docx_tenaga_ahli(dok: Document, data: CVData):
    """Menyusun bagian isi DOCX untuk jenis Tenaga Ahli."""
    _judul_seksi_docx(dok, "A. Data Pribadi")
    _tabel_label_docx(dok, _baris_data_pribadi(data))

    _judul_seksi_docx(dok, "B. Riwayat Pendidikan")
    baris = [
        [b.get("Jenjang", ""), b.get("Jurusan", ""), b.get("Universitas", ""),
         b.get("Kota", ""), b.get("Tahun Lulus", "")]
        for b in data.pendidikan
    ]
    if baris:
        _tabel_data_docx(
            dok, ["Jenjang", "Jurusan", "Universitas", "Kota", "Tahun"],
            baris, [2.0, 4.0, 5.0, 2.8, 2.2],
        )
    else:
        _tulis(dok.add_paragraph(), "Belum diisi.")

    _judul_seksi_docx(dok, "C. Bidang Keahlian")
    p = dok.add_paragraph()
    _tulis(p, "  •  ".join(data.bidang_keahlian) if data.bidang_keahlian else "-")

    _judul_seksi_docx(dok, "D. Sertifikasi Kompetensi")
    baris = [
        [b.get("Nama Sertifikat", ""), b.get("Nomor", ""),
         b.get("Lembaga Penerbit", ""), b.get("Tahun", ""),
         b.get("Masa Berlaku", "")]
        for b in data.sertifikasi
    ]
    if baris:
        _tabel_data_docx(
            dok, ["Nama Sertifikat", "Nomor", "Lembaga", "Tahun", "Berlaku"],
            baris, [4.4, 3.0, 4.0, 2.0, 2.6],
        )
    else:
        _tulis(dok.add_paragraph(), "Tidak ada.")

    _judul_seksi_docx(dok, "E. Pengalaman Kerja")
    baris = []
    for i, b in enumerate(data.pengalaman_kerja, start=1):
        baris.append([
            str(i), b.get("Nama Pekerjaan/Proyek", ""), b.get("Tahun", ""),
            b.get("Jabatan/Posisi", ""), b.get("Instansi Pemberi Kerja", ""),
            b.get("Lokasi", ""),
        ])
    if baris:
        _tabel_data_docx(
            dok, ["No", "Pekerjaan / Proyek", "Tahun", "Jabatan",
                  "Instansi", "Lokasi"],
            baris, [1.0, 4.4, 2.2, 2.8, 3.2, 2.4],
        )
    else:
        _tulis(dok.add_paragraph(), "Belum diisi.")

    if data.publikasi:
        _judul_seksi_docx(dok, "F. Publikasi Ilmiah")
        baris = [
            [str(i), b.get("Judul", ""), b.get("Jurnal/Prosiding", ""),
             b.get("Tahun", "")]
            for i, b in enumerate(data.publikasi, start=1)
        ]
        _tabel_data_docx(
            dok, ["No", "Judul", "Jurnal / Prosiding", "Tahun"],
            baris, [1.0, 7.4, 5.4, 2.2],
        )


def _isi_docx_penyelam(dok: Document, data: CVData):
    """Menyusun bagian isi DOCX untuk jenis Tenaga Spesialis Penyelaman."""
    _judul_seksi_docx(dok, "A. Data Pribadi")
    _tabel_label_docx(dok, _baris_data_pribadi(data))

    _judul_seksi_docx(dok, "B. Sertifikat Selam Utama")
    _tabel_label_docx(dok, [
        ("Jenis Lisensi", data.lisensi_jenis),
        ("Nomor Sertifikat", data.lisensi_nomor),
        ("Level Sertifikasi", data.lisensi_level),
        ("Tanggal Terbit", format_tanggal_indonesia(data.lisensi_terbit)),
        ("Masa Berlaku", format_tanggal_indonesia(data.lisensi_berlaku)
         or "Seumur hidup"),
    ])

    _judul_seksi_docx(dok, "C. Sertifikat Medis Selam (Fit to Dive)")
    _tabel_label_docx(dok, [
        ("Nomor Sertifikat", data.medis_nomor),
        ("Dokter / Klinik Penerbit", data.medis_penerbit),
        ("Tanggal Terbit", format_tanggal_indonesia(data.medis_terbit)),
        ("Berlaku Sampai", format_tanggal_indonesia(data.medis_berlaku)),
    ])

    _judul_seksi_docx(dok, "D. Kompetensi Penyelaman")
    _tabel_label_docx(dok, [
        ("Total Jam Selam Terverifikasi", f"{data.total_jam_selam} jam"),
    ])
    p = dok.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    _tulis(p, "Keahlian Khusus Penyelaman:", tebal=True)
    if data.keahlian_selam:
        for k in data.keahlian_selam:
            pk = dok.add_paragraph()
            pk.paragraph_format.left_indent = Cm(0.8)
            pk.paragraph_format.space_after = Pt(2)
            _tulis(pk, f"\u25aa  {k}")
    else:
        pk = dok.add_paragraph()
        pk.paragraph_format.left_indent = Cm(0.8)
        _tulis(pk, "-")

    _judul_seksi_docx(dok, "E. Pengalaman Penyelaman")
    baris = []
    for i, b in enumerate(data.pengalaman_selam, start=1):
        baris.append([
            str(i), b.get("Lokasi Penyelaman", ""), b.get("Tahun", ""),
            b.get("Jenis Kegiatan", ""), b.get("Kedalaman Maks (m)", ""),
            b.get("Lama Kegiatan", ""), b.get("Pemberi Kerja", ""),
        ])
    if baris:
        _tabel_data_docx(
            dok, ["No", "Lokasi", "Thn", "Jenis Kegiatan", "Maks (m)",
                  "Lama", "Pemberi Kerja"],
            baris, [0.9, 3.2, 1.3, 3.4, 1.7, 2.0, 3.5],
        )
    else:
        _tulis(dok.add_paragraph(), "Belum diisi.")

    _judul_seksi_docx(dok, "F. Pendidikan Formal")
    baris = [
        [b.get("Jenjang", ""), b.get("Jurusan", ""), b.get("Universitas", ""),
         b.get("Kota", ""), b.get("Tahun Lulus", "")]
        for b in data.pendidikan
    ]
    if baris:
        _tabel_data_docx(
            dok, ["Jenjang", "Jurusan", "Universitas", "Kota", "Tahun"],
            baris, [2.0, 4.0, 5.0, 2.8, 2.2],
        )
    else:
        _tulis(dok.add_paragraph(), "Belum diisi.")

    if data.sertifikasi:
        _judul_seksi_docx(dok, "G. Sertifikasi Kompetensi Lainnya")
        baris = [
            [b.get("Nama Sertifikat", ""), b.get("Nomor", ""),
             b.get("Lembaga Penerbit", ""), b.get("Tahun", ""),
             b.get("Masa Berlaku", "")]
            for b in data.sertifikasi
        ]
        _tabel_data_docx(
            dok, ["Nama Sertifikat", "Nomor", "Lembaga", "Tahun", "Berlaku"],
            baris, [4.4, 3.0, 4.0, 2.0, 2.6],
        )
