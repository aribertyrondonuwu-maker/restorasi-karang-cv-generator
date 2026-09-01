"""
st_dinas_luar_builder.py — Mesin pembuat Surat Tugas Dinas Luar.

Menghasilkan dua format keluaran:
  * PDF  (reportlab)  — format utama untuk pengarsipan
  * DOCX (python-docx) — format tambahan untuk penomoran dan tanda tangan basah

Mendukung penugasan perorangan maupun satu tim sekaligus. Surat memuat
klausul pembebasan tugas mengajar dan larangan rangkap bayar sesuai
Diktum KETUJUH SK Dekan.
"""

import io
from dataclasses import dataclass, field
from datetime import date
from typing import List, Dict, Any, Optional

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    Image as RLImage, KeepTogether,
)

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from utils import (
    INSTANSI_KEMENTERIAN, INSTANSI_UNIVERSITAS, INSTANSI_FAKULTAS,
    INSTANSI_ALAMAT, LOGO_PATH, logo_tersedia, DEKAN_NAMA, DEKAN_NIP,
    format_tanggal_indonesia, aman,
)

FONT = "Times-Roman"
FONT_TEBAL = "Times-Bold"
FONT_MIRING = "Times-Italic"
FONT_DOCX = "Times New Roman"

# Margin standar dokumen formal Indonesia
MARGIN_ATAS = 2.0 * cm
MARGIN_BAWAH = 2.5 * cm
MARGIN_KIRI = 3.0 * cm
MARGIN_KANAN = 2.0 * cm
LEBAR_KERJA = 16.0 * cm

WARNA_GARIS = colors.HexColor("#7A8CA0")

NAMA_KEGIATAN_ST = (
    "Swakelola Tipe II: Relokasi dan Replanting/Restorasi Terumbu Karang "
    "di Pelabuhan Perikanan Samudera (PPS) Bitung Tahun Anggaran 2026"
)


# =====================================================================
# STRUKTUR DATA
# =====================================================================

@dataclass
class DataSuratTugas:
    """Data yang dibutuhkan untuk menyusun satu Surat Tugas Dinas Luar."""

    nomor_surat: str = "800/     /UN12.6/TU.00.00/2026"

    # --- Personil yang ditugaskan ---
    # Setiap unsur berupa dict dengan kunci:
    # Nama, NIP, Jabatan, Peran dalam Tim
    personil: List[Dict[str, Any]] = field(default_factory=list)

    unit_kerja: str = (
        "Fakultas Perikanan dan Ilmu Kelautan, Universitas Sam Ratulangi"
    )

    # --- Rincian penugasan ---
    jenis_tugas: str = "Survei Awal Ekologi dan Pemetaan Dasar"
    lokasi: str = "Pelabuhan Perikanan Samudera (PPS) Bitung, Sulawesi Utara"
    tanggal_mulai: Optional[date] = None
    tanggal_selesai: Optional[date] = None
    jumlah_hari: int = 1

    # --- Klausul khusus ---
    klaim_8_oj: bool = False

    # --- Penanda tangan ---
    tempat_ttd: str = "Manado"
    tanggal_ttd: Optional[date] = None

    def satu_orang(self) -> bool:
        """Mengembalikan True bila surat hanya ditujukan untuk satu personil."""
        return len(self.personil) <= 1

    def ada_pembantu_peneliti(self) -> bool:
        """Memeriksa apakah terdapat personil berperan Pembantu Peneliti."""
        return any(
            "pembantu peneliti" in str(p.get("Peran dalam Tim", "")).lower()
            for p in self.personil
        )

    def periode(self) -> str:
        """Menyusun teks periode penugasan dalam Bahasa Indonesia."""
        mulai = format_tanggal_indonesia(self.tanggal_mulai)
        selesai = format_tanggal_indonesia(self.tanggal_selesai)
        if mulai and selesai and mulai != selesai:
            return f"{mulai} s.d. {selesai} ({self.jumlah_hari} hari kerja)"
        if mulai:
            return f"{mulai} ({self.jumlah_hari} hari kerja)"
        return "-"


def _teks_klausul(data: DataSuratTugas) -> str:
    """Menyusun teks klausul sesuai peran dan pilihan klaim honorarium."""
    if data.klaim_8_oj and data.ada_pembantu_peneliti():
        return (
            "Surat tugas ini diterbitkan sebagai syarat pengajuan honorarium "
            "Pembantu Peneliti sebesar 8 (delapan) Orang-Jam (OJ) per hari "
            "kerja. Sebagai konsekuensinya, yang bersangkutan dibebaskan dari "
            "tugas mengajar dan tugas rutin kantor pada tanggal tersebut, "
            "serta tidak menerima uang lauk pauk/tunjangan makan dari unit "
            "kerja dan tidak diperkenankan menerima pembayaran ganda atas "
            "kegiatan yang sama."
        )
    return (
        "Surat tugas ini diterbitkan untuk melegalkan penugasan dinas luar "
        "dan menjadi dasar pengajuan Uang Harian Perjalanan Dinas (UHPD) "
        "sesuai Standar Biaya Masukan (SBM) yang berlaku. Yang bersangkutan "
        "tidak diperkenankan menerima pembayaran ganda atas kegiatan yang sama."
    )


PENUTUP_BAKU = (
    "Demikian surat tugas ini dibuat untuk dilaksanakan dengan penuh "
    "tanggung jawab, dan melaporkan hasil pelaksanaan tugas kepada Dekan "
    "setelah kegiatan selesai."
)


# =====================================================================
# BAGIAN 1 — PEMBANGUN PDF
# =====================================================================

def _gaya_st() -> Dict[str, ParagraphStyle]:
    """Menyusun kumpulan gaya paragraf untuk Surat Tugas berformat PDF."""
    return {
        "kop_kementerian": ParagraphStyle(
            "kop_kementerian", fontName=FONT_TEBAL, fontSize=10, leading=12,
            alignment=TA_CENTER,
        ),
        "kop_besar": ParagraphStyle(
            "kop_besar", fontName=FONT_TEBAL, fontSize=14, leading=16,
            alignment=TA_CENTER,
        ),
        "kop_sedang": ParagraphStyle(
            "kop_sedang", fontName=FONT_TEBAL, fontSize=12, leading=14,
            alignment=TA_CENTER,
        ),
        "kop_alamat": ParagraphStyle(
            "kop_alamat", fontName=FONT, fontSize=8.5, leading=10,
            alignment=TA_CENTER,
        ),
        "judul": ParagraphStyle(
            "judul", fontName=FONT_TEBAL, fontSize=12, leading=15,
            alignment=TA_CENTER,
        ),
        "nomor": ParagraphStyle(
            "nomor", fontName=FONT, fontSize=12, leading=15,
            alignment=TA_CENTER,
        ),
        "isi": ParagraphStyle(
            "isi", fontName=FONT, fontSize=12, leading=20,
            alignment=TA_JUSTIFY,
        ),
        "isi_rapat": ParagraphStyle(
            "isi_rapat", fontName=FONT, fontSize=12, leading=15,
        ),
        "isi_miring": ParagraphStyle(
            "isi_miring", fontName=FONT_MIRING, fontSize=11, leading=17,
            alignment=TA_JUSTIFY,
        ),
        "label": ParagraphStyle(
            "label", fontName=FONT, fontSize=12, leading=15,
        ),
        "tabel_kecil": ParagraphStyle(
            "tabel_kecil", fontName=FONT, fontSize=10.5, leading=13,
        ),
        "tabel_kepala": ParagraphStyle(
            "tabel_kepala", fontName=FONT_TEBAL, fontSize=10.5, leading=13,
            alignment=TA_CENTER,
        ),
        "ttd": ParagraphStyle(
            "ttd", fontName=FONT, fontSize=12, leading=16,
            alignment=TA_CENTER,
        ),
        "ttd_nama": ParagraphStyle(
            "ttd_nama", fontName=FONT_TEBAL, fontSize=12, leading=16,
            alignment=TA_CENTER,
        ),
    }


def _kop_st_pdf(g: Dict[str, ParagraphStyle]) -> List:
    """Menyusun kop surat resmi instansi untuk dokumen PDF."""
    if logo_tersedia():
        try:
            sel_logo = RLImage(LOGO_PATH, width=2.2 * cm, height=2.2 * cm)
        except Exception:
            sel_logo = Paragraph("[LOGO]", g["kop_alamat"])
    else:
        sel_logo = Paragraph("[LOGO]", g["kop_alamat"])

    sel_teks = [
        Paragraph(INSTANSI_KEMENTERIAN, g["kop_kementerian"]),
        Paragraph(INSTANSI_UNIVERSITAS, g["kop_besar"]),
        Paragraph(INSTANSI_FAKULTAS, g["kop_sedang"]),
        Paragraph(INSTANSI_ALAMAT, g["kop_alamat"]),
    ]

    tabel = Table([[sel_logo, sel_teks]], colWidths=[2.6 * cm, 13.4 * cm])
    tabel.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (0, 0), (0, 0), "CENTER"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
    ]))

    garis = Table([[""]], colWidths=[LEBAR_KERJA], rowHeights=[2])
    garis.setStyle(TableStyle([
        ("LINEBELOW", (0, 0), (-1, -1), 2.5, colors.black),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
    ]))

    return [tabel, Spacer(1, 4), garis, Spacer(1, 16)]


def _tabel_personil_pdf(g: Dict[str, ParagraphStyle],
                        data: DataSuratTugas) -> Table:
    """Menyusun tabel daftar personil yang ditugaskan."""
    if data.satu_orang():
        p = data.personil[0] if data.personil else {}
        baris = [
            ("Nama", p.get("Nama", "")),
            ("NIP", p.get("NIP", "")),
            ("Jabatan", p.get("Jabatan", "")),
            ("Unit Kerja", data.unit_kerja),
            ("Peran dalam Kegiatan", p.get("Peran dalam Tim", "")),
        ]
        isi = [
            [Paragraph(label, g["label"]),
             Paragraph(f": {aman(nilai)}", g["label"])]
            for label, nilai in baris
        ]
        tabel = Table(isi, colWidths=[5.0 * cm, 11.0 * cm])
        tabel.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (0, -1), 20),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ]))
        return tabel

    # Lebih dari satu personil: gunakan tabel bernomor
    kepala = ["No", "Nama", "NIP", "Jabatan", "Peran dalam Tim"]
    isi = [[Paragraph(k, g["tabel_kepala"]) for k in kepala]]
    for i, p in enumerate(data.personil, start=1):
        isi.append([
            Paragraph(str(i), g["tabel_kecil"]),
            Paragraph(aman(p.get("Nama", "")), g["tabel_kecil"]),
            Paragraph(aman(p.get("NIP", "")), g["tabel_kecil"]),
            Paragraph(aman(p.get("Jabatan", "")), g["tabel_kecil"]),
            Paragraph(aman(p.get("Peran dalam Tim", "")), g["tabel_kecil"]),
        ])

    tabel = Table(
        isi, colWidths=[1.0 * cm, 4.2 * cm, 3.9 * cm, 3.1 * cm, 3.8 * cm],
        repeatRows=1,
    )
    tabel.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.6, colors.black),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF4")),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return tabel


def generate_st_dinas_luar_pdf(data: DataSuratTugas) -> bytes:
    """Menghasilkan Surat Tugas Dinas Luar berformat PDF."""
    penyangga = io.BytesIO()
    dokumen = SimpleDocTemplate(
        penyangga, pagesize=A4,
        topMargin=MARGIN_ATAS, bottomMargin=MARGIN_BAWAH,
        leftMargin=MARGIN_KIRI, rightMargin=MARGIN_KANAN,
        title=f"Surat Tugas Dinas Luar — {data.nomor_surat}",
        subject=NAMA_KEGIATAN_ST,
    )

    g = _gaya_st()
    e = _kop_st_pdf(g)

    # Judul dan nomor surat
    e.append(Paragraph("<u>SURAT TUGAS DINAS LUAR</u>", g["judul"]))
    e.append(Spacer(1, 2))
    e.append(Paragraph(f"Nomor: {data.nomor_surat}", g["nomor"]))
    e.append(Spacer(1, 16))

    # Kalimat pembuka
    kata_ganti = "kepada" if data.satu_orang() else "kepada nama-nama berikut"
    e.append(Paragraph(
        "Dekan Fakultas Perikanan dan Ilmu Kelautan Universitas Sam "
        f"Ratulangi dengan ini menugaskan {kata_ganti}:", g["isi"]))
    e.append(Spacer(1, 8))
    e.append(_tabel_personil_pdf(g, data))
    e.append(Spacer(1, 12))

    # Rincian tugas
    e.append(Paragraph(
        "Untuk melaksanakan tugas dinas luar dalam rangka kegiatan "
        f"{NAMA_KEGIATAN_ST}, dengan rincian sebagai berikut:", g["isi"]))
    e.append(Spacer(1, 6))

    rincian = [
        ("Jenis Tugas", data.jenis_tugas),
        ("Lokasi Tugas", data.lokasi),
        ("Waktu Pelaksanaan", data.periode()),
    ]
    if not data.satu_orang():
        rincian.insert(0, ("Unit Kerja", data.unit_kerja))

    isi_rincian = [
        [Paragraph(label, g["label"]), Paragraph(f": {aman(nilai)}", g["label"])]
        for label, nilai in rincian
    ]
    tabel_rincian = Table(isi_rincian, colWidths=[5.0 * cm, 11.0 * cm])
    tabel_rincian.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (0, -1), 20),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
    ]))
    e.append(tabel_rincian)
    e.append(Spacer(1, 12))

    # Klausul khusus dan penutup
    e.append(Paragraph(_teks_klausul(data), g["isi_miring"]))
    e.append(Spacer(1, 8))
    e.append(Paragraph(PENUTUP_BAKU, g["isi"]))
    e.append(Spacer(1, 20))

    # Blok tanda tangan Dekan
    tgl_ttd = format_tanggal_indonesia(data.tanggal_ttd or date.today())
    isi_ttd = [
        [Paragraph(f"{aman(data.tempat_ttd)}, {tgl_ttd}", g["ttd"])],
        [Paragraph("Dekan,", g["ttd"])],
        [Spacer(1, 66)],
        [Paragraph(f"<u>{DEKAN_NAMA}</u>", g["ttd_nama"])],
        [Paragraph(f"NIP. {DEKAN_NIP}", g["ttd"])],
    ]
    tabel_ttd = Table(isi_ttd, colWidths=[7.5 * cm], hAlign="RIGHT")
    tabel_ttd.setStyle(TableStyle([
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
    ]))
    e.append(KeepTogether(tabel_ttd))

    dokumen.build(e)
    penyangga.seek(0)
    return penyangga.getvalue()


# =====================================================================
# BAGIAN 2 — PEMBANGUN DOCX
# =====================================================================

def _hapus_tepi_sel(sel):
    """Menghilangkan seluruh garis tepi pada sebuah sel tabel DOCX."""
    tcPr = sel._tc.get_or_add_tcPr()
    tepi = OxmlElement("w:tcBorders")
    for sisi in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{sisi}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tepi.append(el)
    tcPr.append(tepi)


def _atur_lebar_sel(sel, sentimeter: float):
    """Menetapkan lebar sel tabel DOCX dalam satuan sentimeter."""
    tcPr = sel._tc.get_or_add_tcPr()
    for lama in tcPr.findall(qn("w:tcW")):
        tcPr.remove(lama)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(sentimeter * 567)))  # 1 cm = 567 twips
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _tulis(paragraf, teks: str, ukuran: int = 12, tebal: bool = False,
           miring: bool = False, garis_bawah: bool = False):
    """Menambahkan satu potongan teks berformat ke dalam paragraf DOCX."""
    r = paragraf.add_run(teks)
    r.font.name = FONT_DOCX
    r.font.size = Pt(ukuran)
    r.bold = tebal
    r.italic = miring
    r.underline = garis_bawah
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_DOCX)
    rFonts.set(qn("w:cs"), FONT_DOCX)
    return r


def _kop_st_docx(dok: Document):
    """Menyusun kop surat pada bagian header dokumen Word.

    Kop ditempatkan di header agar otomatis muncul kembali apabila surat
    memanjang lebih dari satu halaman.
    """
    bagian = dok.sections[0]
    header = bagian.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.clear()

    tbl = header.add_table(rows=1, cols=2, width=Cm(16.0))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    sel_logo, sel_teks = tbl.rows[0].cells
    _atur_lebar_sel(sel_logo, 2.6)
    _atur_lebar_sel(sel_teks, 13.4)
    _hapus_tepi_sel(sel_logo)
    _hapus_tepi_sel(sel_teks)

    p_logo = sel_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after = Pt(0)
    if logo_tersedia():
        try:
            p_logo.add_run().add_picture(LOGO_PATH, width=Cm(2.1))
        except Exception:
            _tulis(p_logo, "[LOGO]", ukuran=9)
    else:
        _tulis(p_logo, "[LOGO]", ukuran=9)

    baris_kop = [
        (INSTANSI_KEMENTERIAN, 10, True),
        (INSTANSI_UNIVERSITAS, 14, True),
        (INSTANSI_FAKULTAS, 12, True),
        (INSTANSI_ALAMAT, 8, False),
    ]
    pertama = True
    for teks, ukuran, tebal in baris_kop:
        p = sel_teks.paragraphs[0] if pertama else sel_teks.add_paragraph()
        pertama = False
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _tulis(p, teks, ukuran=ukuran, tebal=tebal)

    p_garis = header.add_paragraph()
    p_garis.paragraph_format.space_before = Pt(3)
    p_garis.paragraph_format.space_after = Pt(0)
    pPr = p_garis._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bawah = OxmlElement("w:bottom")
    bawah.set(qn("w:val"), "single")
    bawah.set(qn("w:sz"), "20")
    bawah.set(qn("w:space"), "1")
    bawah.set(qn("w:color"), "000000")
    pBdr.append(bawah)
    pPr.append(pBdr)


def generate_st_dinas_luar(data: DataSuratTugas) -> bytes:
    """Menghasilkan Surat Tugas Dinas Luar berformat DOCX."""
    dok = Document()

    for bagian in dok.sections:
        bagian.top_margin = Cm(2.0)
        bagian.bottom_margin = Cm(2.5)
        bagian.left_margin = Cm(3.0)
        bagian.right_margin = Cm(2.0)
        bagian.header_distance = Cm(1.0)

    gaya = dok.styles["Normal"]
    gaya.font.name = FONT_DOCX
    gaya.font.size = Pt(12)

    _kop_st_docx(dok)

    # Jarak setelah kop
    p_jarak = dok.add_paragraph()
    p_jarak.paragraph_format.space_after = Pt(6)

    # Judul dan nomor
    p_judul = dok.add_paragraph()
    p_judul.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_judul.paragraph_format.space_after = Pt(2)
    _tulis(p_judul, "SURAT TUGAS DINAS LUAR", tebal=True, garis_bawah=True)

    p_nomor = dok.add_paragraph()
    p_nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nomor.paragraph_format.space_after = Pt(14)
    _tulis(p_nomor, f"Nomor: {data.nomor_surat}")

    # Kalimat pembuka
    kata_ganti = "kepada" if data.satu_orang() else "kepada nama-nama berikut"
    p_buka = dok.add_paragraph()
    p_buka.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_buka.paragraph_format.line_spacing = 1.5
    p_buka.paragraph_format.space_after = Pt(6)
    _tulis(p_buka,
           "Dekan Fakultas Perikanan dan Ilmu Kelautan Universitas Sam "
           f"Ratulangi dengan ini menugaskan {kata_ganti}:")

    # Daftar personil
    if data.satu_orang():
        p = data.personil[0] if data.personil else {}
        baris = [
            ("Nama", p.get("Nama", "")),
            ("NIP", p.get("NIP", "")),
            ("Jabatan", p.get("Jabatan", "")),
            ("Unit Kerja", data.unit_kerja),
            ("Peran dalam Kegiatan", p.get("Peran dalam Tim", "")),
        ]
        tbl = dok.add_table(rows=len(baris), cols=2)
        for i, (label, nilai) in enumerate(baris):
            sel_l, sel_r = tbl.cell(i, 0), tbl.cell(i, 1)
            _atur_lebar_sel(sel_l, 5.0)
            _atur_lebar_sel(sel_r, 11.0)
            _hapus_tepi_sel(sel_l)
            _hapus_tepi_sel(sel_r)
            p_l = sel_l.paragraphs[0]
            p_l.paragraph_format.left_indent = Cm(1.0)
            _tulis(p_l, label)
            _tulis(sel_r.paragraphs[0], f": {aman(nilai)}")
    else:
        kepala = ["No", "Nama", "NIP", "Jabatan", "Peran dalam Tim"]
        lebar = [1.0, 4.2, 3.9, 3.1, 3.8]
        tbl = dok.add_table(rows=len(data.personil) + 1, cols=len(kepala))
        tbl.style = "Table Grid"
        for j, judul in enumerate(kepala):
            sel = tbl.cell(0, j)
            _atur_lebar_sel(sel, lebar[j])
            p = sel.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _tulis(p, judul, ukuran=11, tebal=True)
        for i, orang in enumerate(data.personil, start=1):
            nilai = [
                str(i), orang.get("Nama", ""), orang.get("NIP", ""),
                orang.get("Jabatan", ""), orang.get("Peran dalam Tim", ""),
            ]
            for j, v in enumerate(nilai):
                sel = tbl.cell(i, j)
                _atur_lebar_sel(sel, lebar[j])
                _tulis(sel.paragraphs[0], aman(v), ukuran=11)

    # Rincian tugas
    p_kosong = dok.add_paragraph()
    p_kosong.paragraph_format.space_after = Pt(4)

    p_rinci = dok.add_paragraph()
    p_rinci.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_rinci.paragraph_format.line_spacing = 1.5
    p_rinci.paragraph_format.space_after = Pt(6)
    _tulis(p_rinci,
           "Untuk melaksanakan tugas dinas luar dalam rangka kegiatan "
           f"{NAMA_KEGIATAN_ST}, dengan rincian sebagai berikut:")

    rincian = [
        ("Jenis Tugas", data.jenis_tugas),
        ("Lokasi Tugas", data.lokasi),
        ("Waktu Pelaksanaan", data.periode()),
    ]
    tbl_r = dok.add_table(rows=len(rincian), cols=2)
    for i, (label, nilai) in enumerate(rincian):
        sel_l, sel_r = tbl_r.cell(i, 0), tbl_r.cell(i, 1)
        _atur_lebar_sel(sel_l, 5.0)
        _atur_lebar_sel(sel_r, 11.0)
        _hapus_tepi_sel(sel_l)
        _hapus_tepi_sel(sel_r)
        p_l = sel_l.paragraphs[0]
        p_l.paragraph_format.left_indent = Cm(1.0)
        _tulis(p_l, label)
        _tulis(sel_r.paragraphs[0], f": {aman(nilai)}")

    # Klausul dan penutup
    p_kosong2 = dok.add_paragraph()
    p_kosong2.paragraph_format.space_after = Pt(4)

    p_klausul = dok.add_paragraph()
    p_klausul.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_klausul.paragraph_format.line_spacing = 1.5
    p_klausul.paragraph_format.space_after = Pt(6)
    _tulis(p_klausul, _teks_klausul(data), ukuran=11, miring=True)

    p_tutup = dok.add_paragraph()
    p_tutup.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_tutup.paragraph_format.line_spacing = 1.5
    p_tutup.paragraph_format.space_after = Pt(14)
    _tulis(p_tutup, PENUTUP_BAKU)

    # Blok tanda tangan Dekan
    tgl_ttd = format_tanggal_indonesia(data.tanggal_ttd or date.today())
    tbl_ttd = dok.add_table(rows=1, cols=2)
    tbl_ttd.alignment = WD_TABLE_ALIGNMENT.RIGHT
    sel_kosong, sel_ttd = tbl_ttd.rows[0].cells
    _atur_lebar_sel(sel_kosong, 8.5)
    _atur_lebar_sel(sel_ttd, 7.5)
    _hapus_tepi_sel(sel_kosong)
    _hapus_tepi_sel(sel_ttd)

    def _baris_ttd(teks, tebal=False, garis_bawah=False, pertama=False):
        p = sel_ttd.paragraphs[0] if pertama else sel_ttd.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        _tulis(p, teks, tebal=tebal, garis_bawah=garis_bawah)

    _baris_ttd(f"{aman(data.tempat_ttd)}, {tgl_ttd}", pertama=True)
    _baris_ttd("Dekan,")
    for _ in range(4):
        _baris_ttd("")
    _baris_ttd(DEKAN_NAMA, tebal=True, garis_bawah=True)
    _baris_ttd(f"NIP. {DEKAN_NIP}")

    keluaran = io.BytesIO()
    dok.save(keluaran)
    keluaran.seek(0)
    return keluaran.getvalue()
